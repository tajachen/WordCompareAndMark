# MyProject - Example Project under GPL v3 License
#
# Copyright (C) 2024 Your Name or Your Organization's Name
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <http://www.gnu.org/licenses/>.

import datetime
import tkinter as tk
from tkinter import filedialog, messagebox

from docx import Document
from docx.shared import RGBColor


# 1.暂时只处理docx文档，后面考虑也引入doc的库

def select_file1(button, label):
    filename = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if filename:
        # print("Selected file:", filename)
        label.config(text=filename)  # 如果已选择文件，显示文件名完整路径
    else:
        label.config(text="")  # 如果未选择文件，则标签不显示文件名和后缀
    return filename


def select_file2(button, label):
    filename = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if filename:
        # print("Selected file:", filename)
        label.config(text=filename)  # 如果已选择文件，显示文件名完整路径
    else:
        label.config(text="")  # 如果未选择文件，则标签不显示文件名和后缀
    return filename


def compare_files(file1, file2):
    # 判断2个文件是否都已被选择
    if not file1 or not file2:
        messagebox.showwarning("Warning", "请先选择两个文件")
        return

    # 判断2个文件是否一致
    if file1 == file2:
        messagebox.showwarning("Warning", "两个文件一致，请重新选择")
        return

    # 尝试打开文件，如果文件已被其他程序打开，将会抛出异常
    try:
        with open(file1, 'r+') as f:
            pass
    except IOError:
        messagebox.showwarning("Warning", "文件1 已被其他程序打开")
        return

    try:
        with open(file2, 'r+') as f:
            pass
    except IOError:
        messagebox.showwarning("Warning", "文件2 已被其他程序打开")
        return

    try:
        doc1 = Document(file1)
        doc2 = Document(file2)

        # 将文件1和文件2的所有字体设置为黑色
        for doc in [doc1, doc2]:
            for para in doc.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为黑色
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为黑色

        text1 = " ".join([para.text for para in doc1.paragraphs])
        text2 = " ".join([para.text for para in doc2.paragraphs])
        # text1 = " ".join(
        #     [para.text for para in doc1.paragraphs] + [cell.text for table in doc1.tables for row in table.rows for cell
        #                                                in row.cells])
        # text2 = " ".join(
        #     [para.text for para in doc2.paragraphs] + [cell.text for table in doc2.tables for row in table.rows for cell
        #                                                in row.cells])

        time_begin = "查重开始时间：" + datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # 查找连续10个字符一致的内容
        for i in range(len(text1) - 9):
            substring1 = text1[i:i + 10]
            if substring1 in text2:
                # 对一致的内容进行红色字体标记
                for para in doc1.paragraphs:
                    for run in para.runs:
                        if substring1 in run.text:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # 设置字体颜色为红色
                # for table in doc1.tables:
                #     for row in table.rows:
                #         for cell in row.cells:
                #             for para in cell.paragraphs:
                #                 for run in para.runs:
                #                     if substring in run.text:
                #                         run.font.color.rgb = RGBColor(255, 0, 0)  # 设置字体颜色为红色

        time_1_end = "文件1查重完毕时间：" + datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # 查找连续10个字符一致的内容
        for i in range(len(text2) - 9):
            substring2 = text2[i:i + 10]
            if substring2 in text1:
                # 对一致的内容进行红色字体标记
                for para in doc2.paragraphs:
                    for run in para.runs:
                        if substring2 in run.text:
                            run.font.color.rgb = RGBColor(255, 0, 0)  # 设置字体颜色为红色
                # for table in doc2.tables:
                #     for row in table.rows:
                #         for cell in row.cells:
                #             for para in cell.paragraphs:
                #                 for run in para.runs:
                #                     if substring in run.text:
                #                         run.font.color.rgb = RGBColor(255, 0, 0)  # 设置字体颜色为红色

        time_2_end = "文件2查重完毕时间：" + datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # 保存文档
        doc1.save(file1)
        doc2.save(file2)

        # 显示处理成功的弹窗
        messagebox.showinfo("Info", "已查重完毕。\n"
                                    "2个文件的[红色字体]为疑似重复项，请人工复查。\n" +
                            time_begin + "\n" +
                            time_1_end + "\n" +
                            time_2_end)

    except BaseException as e:
        messagebox.showwarning("Warning", "异常信息为：" + str(e))
        return


if __name__ == '__main__':
    # 创建主窗口
    root = tk.Tk()
    # 设置窗口标题为
    root.title("Word文档查重并标记_V1.0.2")
    # 设置窗口大小
    root.geometry("600x600")

    # 创建一个Frame作为备注的容器
    frame_ps = tk.Frame(root, bd=1, relief='sunken')
    frame_ps.pack(anchor='nw', pady=10, padx=10)  # 设置在窗口左上方，上下左右间隔为20

    # Frame中创建一个只读的文本区域
    text = tk.Text(frame_ps, height=28, width=75)
    text.insert(tk.END, "操作方式：\n"
                        "\n"
                        "1.分别选择2个docx文件（暂不支持doc文件），因为是直接对原文档编辑，所以需要提前做好备份；\n"
                        "2.确保2个docx文件是未打开状态（工具也会判断）；\n"
                        "3.确保2个docx文件全文都是黑色字体（工具运行时也会先将2个文件的所有字体强行设置为黑色字体）；\n"
                        "4.点击[查重并标记保存]按钮，开始查重并标记（2个文档各100页左右，预计运行2分钟，请耐心等待）；\n"
                        "5.弹窗提示成功后，2个文件的疑似重复项将被标记为[红色字体]，再由人工复查。\n"
                        "\n"
                        "\n"
                        "\n"
                        "\n"
                        "---------------------------------------------------------------------------\n"
                        "\n"
                        "处理机制：\n"
                        "1.[文件1]的每连续10个字符，是否在[文件2]中出现，[文件1]重复的文本所在段落被设置为红色字体；\n"
                        "2.反向，[文件2]的每连续10个字符，是否在[文件1]中出现，[文件2]重复的文本所在段落被设置为红色字体。\n"
                        "\n"
                        "---------------------------------------------------------------------------\n"
                        "\n"
                        "版本更新：V1.0.1\n"
                        "1.根据开发者查重经验进行基础编写，现仅对文本进行查重。\n"
                        "版本更新：V1.0.2\n"
                        "1.强制设置黑色字体也包含了表格内容。\n"
                        "2.表格内容暂不参与查重，没必要。\n"
                )
    text.configure(state='disabled')  # 设置为只读

    # Frame中创建一个滚动条
    scrollbar = tk.Scrollbar(frame_ps, command=text.yview)
    text['yscrollcommand'] = scrollbar.set

    # Frame中平行放置文本区域和滚动条
    text.pack(side=tk.LEFT, fill=tk.Y, pady=10, padx=10)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    # 创建一个Frame作为按钮的容器
    frame = tk.Frame(root, bd=1, relief='sunken')
    frame.pack(anchor='nw', pady=10, padx=20)  # 设置在窗口左上方，上下左右间隔为20

    # 创建两个选择文件的按钮
    button1 = tk.Button(frame, text="请选择文件1", command=lambda: select_file1(button1, label1))
    label1 = tk.Label(frame, text="")
    button2 = tk.Button(frame, text="请选择文件2", command=lambda: select_file2(button2, label2))
    label2 = tk.Label(frame, text="")

    # 创建一个查重按钮
    compare_button = tk.Button(root, text="查重并标记保存",
                               command=lambda: compare_files(label1.cget("text"), label2.cget("text")))
    # 平行放置按钮和标签，上下间隔为20
    button1.grid(row=0, column=0, pady=10, padx=10)
    label1.grid(row=0, column=1, pady=10, padx=10)
    button2.grid(row=1, column=0, pady=10, padx=10)
    label2.grid(row=1, column=1, pady=10, padx=10)
    compare_button.pack(side=tk.TOP, pady=10)

    # 运行主循环
    root.mainloop()
